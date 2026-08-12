"""
core/utils/file_parser.py - 文件解析工具
支持 PDF (.pdf)、Word (.docx)、纯文本 (.txt) 三种常见简历/JD文件格式的解析。
"""

from pathlib import Path
from typing import Optional

from loguru import logger


# ------------------------------------------------------------
# 单文件解析函数
# ------------------------------------------------------------

def parse_pdf(file_path: str) -> str:
    """
    解析 PDF 文件，返回全部文本内容。

    Args:
        file_path: PDF 文件路径

    Returns:
        提取出的纯文本字符串
    """
    try:
        from PyPDF2 import PdfReader

        text_parts: list[str] = []
        reader = PdfReader(file_path)
        for page_num, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
            else:
                logger.debug(f"PDF 第 {page_num} 页无文本内容")

        full_text = "\n".join(text_parts)
        logger.info(f"PDF 解析完成: {file_path} | 共 {len(reader.pages)} 页 | {len(full_text)} 字符")
        return full_text

    except ImportError:
        logger.error("PyPDF2 未安装，请执行: pip install PyPDF2")
        raise
    except Exception as e:
        logger.error(f"PDF 解析失败 [{file_path}]: {e}")
        raise


def parse_docx(file_path: str) -> str:
    """
    解析 Word (.docx) 文件，返回全部文本内容。

    Args:
        file_path: DOCX 文件路径

    Returns:
        提取出的纯文本字符串
    """
    try:
        from docx import Document

        doc = Document(file_path)
        text_parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        # 同时提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))

        full_text = "\n".join(text_parts)
        logger.info(f"DOCX 解析完成: {file_path} | {len(full_text)} 字符")
        return full_text

    except ImportError:
        logger.error("python-docx 未安装，请执行: pip install python-docx")
        raise
    except Exception as e:
        logger.error(f"DOCX 解析失败 [{file_path}]: {e}")
        raise


def parse_txt(file_path: str) -> str:
    """
    解析纯文本文件，自动尝试常见编码。

    Args:
        file_path: TXT 文件路径

    Returns:
        文件全部文本内容
    """
    encodings_to_try = ["utf-8", "gbk", "gb2312", "latin-1"]

    for enc in encodings_to_try:
        try:
            with open(file_path, "r", encoding=enc) as f:
                text = f.read()
            logger.info(f"TXT 解析完成: {file_path} | 编码={enc} | {len(text)} 字符")
            return text
        except (UnicodeDecodeError, UnicodeError):
            continue

    # 所有编码均失败，最后使用 utf-8 with errors='replace'
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    logger.warning(f"TXT 解析 (fallback): {file_path} | 使用 utf-8-replace 模式")
    return text


# ------------------------------------------------------------
# 统一解析入口
# ------------------------------------------------------------

# 支持的文件扩展名
SUPPORTED_EXTENSIONS = {
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".txt": parse_txt,
}


def parse_file(file_path: str) -> str:
    """
    根据文件扩展名自动选择解析器，解析文件内容。

    Args:
        file_path: 文件路径

    Returns:
        提取出的文本内容

    Raises:
        ValueError: 不支持的文件格式
        FileNotFoundError: 文件不存在
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"不支持的文件格式: {ext}，目前支持: {list(SUPPORTED_EXTENSIONS.keys())}"
        )

    parser = SUPPORTED_EXTENSIONS[ext]
    logger.debug(f"使用解析器 [{parser.__name__}] 解析文件: {file_path}")
    return parser(file_path)


# ------------------------------------------------------------
# 批量解析
# ------------------------------------------------------------

def parse_files(file_paths: list[str]) -> dict[str, str]:
    """
    批量解析多个文件。

    Args:
        file_paths: 文件路径列表

    Returns:
        {文件路径: 解析文本} 字典，解析失败的文件不会出现在结果中
    """
    results: dict[str, str] = {}
    for fp in file_paths:
        try:
            results[fp] = parse_file(fp)
        except Exception as e:
            logger.error(f"批量解析跳过 [{fp}]: {e}")
    logger.info(f"批量解析完成: 成功 {len(results)}/{len(file_paths)} 个文件")
    return results
