"""
core/tools/file_export.py - 文件导出模块

功能:
    1. 优化简历导出为 .docx（格式化排版）
    2. 面试报告导出为 .txt
    3. 匹配分析报告导出为 .txt
    4. 对比报告导出为 .txt
"""

import os
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from loguru import logger

from config.settings import settings, PROJECT_ROOT


# ============================================================
# 默认输出目录
# ============================================================

_EXPORT_DIR: Optional[Path] = None


def get_export_dir() -> Path:
    """获取导出目录，确保目录存在。"""
    global _EXPORT_DIR
    if _EXPORT_DIR is None:
        _EXPORT_DIR = PROJECT_ROOT / "exports"
    _EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    return _EXPORT_DIR


def set_export_dir(path: str) -> None:
    """设置自定义导出目录。"""
    global _EXPORT_DIR
    _EXPORT_DIR = Path(path)
    _EXPORT_DIR.mkdir(parents=True, exist_ok=True)


# ============================================================
# Docx 导出
# ============================================================

def _apply_docx_style(doc: Document) -> None:
    """为文档应用统一的中文排版样式。"""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "微软雅黑"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(4)

    # 设置默认中文字体（通过 run-level 在写入时处理）
    # 这里确保标题样式
    for level in range(1, 4):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.name = "微软雅黑"
        heading_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def _add_docx_heading(doc: Document, text: str, level: int = 1) -> None:
    """添加标题并设置格式。"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "微软雅黑"
        if level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(12)
    return heading


def _add_docx_paragraph(doc: Document, text: str, bold: bool = False,
                        size: Optional[int] = None, alignment=None) -> None:
    """添加段落并设置格式。"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "微软雅黑"
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    if alignment is not None:
        para.alignment = alignment
    para.paragraph_format.line_spacing = 1.5
    return para


def _add_docx_bullet(doc: Document, text: str) -> None:
    """添加项目符号列表项。"""
    para = doc.add_paragraph(style="List Bullet")
    # 清除默认 runs 并添加自定义 run
    for r in para.runs:
        r.text = ""
    run = para.add_run(text)
    run.font.name = "微软雅黑"
    run.font.size = Pt(11)


def export_resume_to_docx(
    optimized_resume: str,
    job_title: str = "应聘简历",
    company_name: str = "",
    applicant_name: str = "",
    match_scores: Optional[Dict[str, Any]] = None,
    output_path: Optional[str] = None,
) -> str:
    """
    将优化后的简历导出为格式化的 .docx 文件。

    Args:
        optimized_resume: 优化后的简历文本（Markdown 或纯文本）
        job_title:        应聘职位名称
        company_name:     目标公司名称
        applicant_name:   求职者姓名
        match_scores:     匹配分数字典（可选），用于附加匹配分析页
        output_path:      输出文件路径，为 None 时自动生成

    Returns:
        生成的文件绝对路径
    """
    if output_path is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_job = job_title.replace(" ", "_").replace("/", "-")
        filename = f"简历优化_{safe_job}_{timestamp}.docx"
        output_path = str(get_export_dir() / filename)

    doc = Document()
    _apply_docx_style(doc)

    # === 封面区域 ===
    _add_docx_heading(doc, "优化简历", level=1)

    # 基本信息行
    info_lines = []
    if applicant_name:
        info_lines.append(f"求职者: {applicant_name}")
    info_lines.append(f"目标职位: {job_title}")
    if company_name:
        info_lines.append(f"目标公司: {company_name}")
    info_lines.append(f"生成日期: {datetime.now().strftime('%Y年%m月%d日')}")

    for line in info_lines:
        _add_docx_paragraph(doc, line, size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_paragraph()  # 空行分隔

    # === 简历正文 ===
    _add_docx_heading(doc, "简历正文", level=2)

    # 按段落解析简历内容
    sections = optimized_resume.strip().split("\n\n")
    for section in sections:
        section = section.strip()
        if not section:
            continue

        lines = section.split("\n")
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue

            # 检测 Markdown 标题
            if line.startswith("### "):
                _add_docx_heading(doc, line[4:], level=3)
            elif line.startswith("## "):
                _add_docx_heading(doc, line[3:], level=2)
            elif line.startswith("# "):
                _add_docx_heading(doc, line[2:], level=2)
            # 检测 Markdown 列表
            elif line.startswith("- ") or line.startswith("* ") or line.startswith("• "):
                _add_docx_bullet(doc, line[2:])
            elif line.startswith("**") and "**" in line[2:]:
                # 粗体段落
                clean = line.replace("**", "")
                _add_docx_paragraph(doc, clean, bold=True)
            else:
                _add_docx_paragraph(doc, line)

    # === 匹配分析页（可选）===
    if match_scores:
        doc.add_page_break()
        _add_docx_heading(doc, "匹配分析", level=2)

        overall = match_scores.get("overall", {})
        _add_docx_paragraph(doc,
            f"综合得分: {overall.get('after', overall.get('score', 'N/A'))}/100",
            bold=True, size=14)

        dims = [
            ("技能匹配度", "skill_match"),
            ("经验匹配度", "experience_match"),
            ("学历匹配度", "education_match"),
            ("JD关键词覆盖", "jd_keyword_coverage"),
        ]
        for label, key in dims:
            val = match_scores.get(key, {})
            score_val = val.get("after", val) if isinstance(val, dict) else val
            _add_docx_paragraph(doc, f"{label}: {score_val}/100")

        analysis = match_scores.get("analysis_before", "") or match_scores.get("analysis", "")
        if analysis:
            _add_docx_heading(doc, "分析说明", level=3)
            _add_docx_paragraph(doc, analysis)

    doc.save(output_path)
    logger.info(f"[file_export] 简历已导出: {output_path}")
    return str(Path(output_path).resolve())


# ============================================================
# TXT 导出
# ============================================================

def _ensure_txt_output_path(output_path: Optional[str], prefix: str) -> str:
    """为 txt 文件生成输出路径。"""
    if output_path is not None:
        return output_path
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{prefix}_{timestamp}.txt"
    return str(get_export_dir() / filename)


def export_interview_report_to_txt(
    qa_pairs: List[Dict[str, Any]],
    job_title: str = "面试复盘",
    applicant_name: str = "",
    overall_score: Optional[float] = None,
    output_path: Optional[str] = None,
) -> str:
    """
    将面试问答记录导出为 .txt 面试报告。

    Args:
        qa_pairs:       面试问答列表，每个元素为 {question, answer, score, feedback}
        job_title:      应聘职位
        applicant_name: 求职者姓名
        overall_score:  面试总评分（可选）
        output_path:    输出路径

    Returns:
        生成的文件绝对路径
    """
    output_path = _ensure_txt_output_path(output_path, "面试报告")

    lines = [
        "=" * 60,
        f"  面试复盘报告",
        "=" * 60,
        "",
        f"  职位: {job_title}",
        f"  求职者: {applicant_name or '(未填写)'}",
        f"  生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
    ]

    if overall_score is not None:
        lines.append(f"  面试总评: {overall_score}/100")
        lines.append("")

    lines.append("-" * 60)
    lines.append(f"  共 {len(qa_pairs)} 道面试题")
    lines.append("-" * 60)
    lines.append("")

    for i, qa in enumerate(qa_pairs, 1):
        lines.append(f"【第 {i} 题】")
        lines.append(f"  问题: {qa.get('question', '')}")
        lines.append("")
        lines.append(f"  回答: {qa.get('answer', '')}")
        lines.append("")

        score = qa.get("score")
        if score is not None:
            lines.append(f"  得分: {score}")
        feedback = qa.get("feedback", "")
        if feedback:
            lines.append(f"  点评: {feedback}")
        lines.append("")

    lines.append("-" * 60)
    lines.append("  面试建议:")
    lines.append("  1. 针对薄弱环节加强练习")
    lines.append("  2. 准备 2-3 个工作相关的成功案例（STAR 法则）")
    lines.append("  3. 模拟面试时注意语速和表达的清晰度")
    lines.append("")
    lines.append("=" * 60)
    lines.append("  ResumeAgent · AI 多智能体求职助手 自动生成")
    lines.append("=" * 60)

    content = "\n".join(lines)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)

    logger.info(f"[file_export] 面试报告已导出: {output_path}")
    return str(Path(output_path).resolve())


def export_match_report_to_txt(
    report_text: str,
    job_title: str = "JD匹配分析",
    output_path: Optional[str] = None,
) -> str:
    """
    将匹配分析报告文本导出为 .txt 文件。

    Args:
        report_text: 报告文本内容
        job_title:   职位名称
        output_path: 输出路径

    Returns:
        生成的文件绝对路径
    """
    output_path = _ensure_txt_output_path(output_path, "匹配报告")

    lines = [
        f"  职位: {job_title}",
        f"  生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        report_text,
        "",
        "=" * 60,
        "  ResumeAgent · AI 多智能体求职助手 自动生成",
        "=" * 60,
    ]
    content = "\n".join(lines)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)

    logger.info(f"[file_export] 匹配报告已导出: {output_path}")
    return str(Path(output_path).resolve())


def export_compare_report_to_txt(
    compare_text: str,
    job_title: str = "简历优化对比",
    output_path: Optional[str] = None,
) -> str:
    """
    将优化前后对比报告导出为 .txt 文件。

    Args:
        compare_text: 对比报告文本
        job_title:    职位名称
        output_path:  输出路径

    Returns:
        生成的文件绝对路径
    """
    output_path = _ensure_txt_output_path(output_path, "对比报告")

    lines = [
        f"  职位: {job_title}",
        f"  生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        compare_text,
        "",
        "=" * 60,
        "  ResumeAgent · AI 多智能体求职助手 自动生成",
        "=" * 60,
    ]
    content = "\n".join(lines)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)

    logger.info(f"[file_export] 对比报告已导出: {output_path}")
    return str(Path(output_path).resolve())
